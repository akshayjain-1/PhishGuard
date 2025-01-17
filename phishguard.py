# pylint: disable=W1203
# pylint: disable=C0301
"""
Program to detect phishing emails
"""
from asyncio import sleep
from datetime import datetime
import email
from email.utils import parseaddr
import hashlib
import os
import re
import sys
import logging
from email.header import decode_header
from email import policy
from email.parser import BytesParser
from io import BytesIO
from docx import Document
from bs4 import BeautifulSoup
import dns.resolver
import dkim
import vt
from pyzbar.pyzbar import decode
from PIL import Image
import fitz
import checkdmarc

# Constants
VIRUSTOTAL_API_KEY = os.environ.get('VIRUSTOTAL_API_KEY')  # Replace with your actual API key
if not VIRUSTOTAL_API_KEY:
    logging.critical("VIRUSTOTAL_API_KEY is not set. Exiting.")
    sys.exit(1)

RECENT_THRESHOLD_DAYS = 30
PHISHING_KEYWORDS = ['invoice', 'order', 'purchase', 'confirm', 'account',
                     'password', 'urgent', 'payment', 'subscription', 'update your account', 
                     'verify your account', 'suspend']
KNOWN_BRANDS = ['mcafee', 'paypal', 'amazon', 'microsoft', 'geek squad']
PATTERNS = {
    "Sextortion": r'\b(sextortion|blackmail|explicit|photographs)\b',
    "Bank Request": r'\b(update|confirm|verify|access|suspend)\b.*\b(account|bank)\b'
}


# Configure logging
log_file_path = 'email_analyzer.log'
if os.path.isfile(log_file_path):
    os.remove(log_file_path)

logging.basicConfig(
    filename='email_analyzer.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

def how_to_save_email() -> None:
    """
    Function to guide the user to save email as .eml if needed
    """
    instructions = """
    How to save an email as a .eml file:

    - **Gmail**: Open the email > Click 'More' (three dots) nect to Reply button  > 'Download message'
    - **Outlook**: Open the email > 'File' > 'Save As' > Save as .eml file
    - **Thunderbird**: Right-click the email > 'Save As' > Save as .eml file
    - **Apple Mail**: Open the email > 'File' > 'Save As' > Choose the destination and ensure the file format is set to .eml

    """
    print(instructions)

def parse_email(file_path: str):
    """
    Function to parse email
    """
    try:
        with open(file_path, 'rb') as f:
            msg = BytesParser(policy=policy.default).parse(f)
        return msg
    except FileNotFoundError:
        logging.error(f"File not found: {file_path}")  # pylint:disable=logging-fstring-interpolation
        print(f"Error: File '{file_path}' not found.")
        return None


# Validate sender's domain using SPF, DKIM, and DMARC
def _check_spf(domain: str) -> str:
    """
    Function to check spf
    """
    try:
        answers = checkdmarc.check_spf(domain)
        if 'v=spf1' in answers['record'] or answers['valid']:
            logging.info(f"SPF record is valid.")
            return "SPF record is valid."
        if not answers['valid']:
            logging.info(f"No SPF record found.")
            return "No SPF record found."
    except Exception as e:  # pylint: disable=W0718
        logging.error(f"Error checking SPF for domain '{domain}': {e}")
    return "SPF validation check error."

def _check_dmarc(domain: str) -> str:
    """
    Function to check dmarc
    """
    try:
        answers = checkdmarc.check_dmarc(domain)
        if 'v=spf1' in answers['record'] or answers['valid']:
            logging.info(f"DMARC record is valid.")
            return "DMARC record is valid."
        if not answers['valid']:
            logging.info(f"No DMARC record found")
            return "No DMARC record found."
    except Exception as e:  # pylint: disable=W0718
        logging.error(f"Error checking DMARC for domain '{domain}': {e}")
    return "DMARC validation check error."

def _validate_dkim(msg) -> str:
    """
    Validate DKIM signature of the email.
    """
    raw_email = msg.as_bytes()
    result = dkim.verify(raw_email)
    if result:
        logging.info("DKIM signature is valid.")
        return "DKIM signature is valid."
    logging.info("DKIM signature is invalid.")
    return "DKIM validation failed."

def validate_sender_address(email_from, msg):
    """
    Perform validation of sender's address via spf, dkim and dmarc check
    """
    from_address = email_from.split('<')[-1].strip().rstrip('>')
    domain = from_address.split('@')[-1]
    spf_status = _check_spf(domain)
    dmarc_status = _check_dmarc(domain)
    dkim_status = _validate_dkim(msg)
    return spf_status, dmarc_status, dkim_status

def extract_email_headers(msg):
    """
    Extracts and decodes email headers securely.
    """
    headers = {}
    for key, value in msg.items():
        decoded_value = decode_header(value)[0][0]
        headers[key] = (
            decoded_value if isinstance(decoded_value, str) else decoded_value.decode(errors="ignore")
        )
    logging.info("Extracted email headers.")
    return headers

def analyze_sender_address(headers):
    """
    Analyzes the sender address for common phishing indicators.
    """
    anomalies = ''
    from_address = headers.get('From', '')
    if not from_address or '@' not in from_address:
        return "Invalid sender address."

    from_address = from_address.split('<')[-1].strip().rstrip('>')
    
    domain = from_address.split('@')[-1]
    
    result = checkdmarc.check_mx(domain)

    if result['hosts']:
        anomalies = "Sender domain exists and has a valid MX record."
    else:
        anomalies = "Potentially suspicious sender domain as it doesn't have a MX record."
    
    logging.info(f"Sender address: {anomalies}")
    return anomalies

def analyze_headers_for_anomalies(headers):
    """
    Checks headers for anomalies such as mismatched 'Reply-To' and 'From' fields.
    """
    anomalies = []
    if headers.get('Reply-To') and headers['Reply-To'] != headers.get('From'):
        anomalies.append("Mismatch between 'Reply-To' and 'From' fields.")
    
    if not headers.get('Reply-To'):
        anomalies.append("Missing 'Reply-To' field.")

    if 'X-Priority' in headers and headers['X-Priority'] in ['1', 'High']:
        anomalies.append("Email marked as high priority, could indicate urgency-based phishing.")

    logging.info(f"Header Anomalies: {anomalies}")
    return anomalies

def extract_urls(msg):
    """
    Extracts URLs from the email body.
    """
    urls = []
    for part in msg.walk():
        if part.get_content_type() == 'text/html':
            soup = BeautifulSoup(part.get_payload(decode=True), 'html.parser')
            urls.extend([link.get('href') for link in soup.find_all('a', href=True)])
    logging.info(f"Extracted URLs: {urls}")
    return urls

def check_url_with_virustotal(urls):
    """
    Checks a list of URLs with VirusTotal and returns the analysis results.
    """
    results = {}
    
    with vt.Client(VIRUSTOTAL_API_KEY) as client:
        for url in urls:
            try:
                analysis = client.get_object(f"/urls/{vt.url_id(url)}")
                stats = analysis.last_analysis_stats
                malicious_count = stats.get('malicious', 0)
                # Get the domain creation date
                creation_date = _get_domain_creation_date(analysis)
                # Check if the domain was recently created
                if creation_date and _is_recent_domain(creation_date, RECENT_THRESHOLD_DAYS):
                    verdict = f"URL {url} is suspicious: Domain is newly created - {creation_date}"
                elif malicious_count > 0:
                    verdict = f"URL '{url}' is flagged as malicious."
                else:
                    verdict = f"URL '{url}' appears safe."
                # Store the result
                results[url] = verdict
                logging.info(f"Checked URL: {url}, Verdict: {verdict}")
            except Exception as e:  # pylint: disable=W0718
                logging.exception(f"Unexpected error for URL '{url}': {e}")
                results[url] = f"Unexpected error occurred. {e}"

    return results

def _get_domain_creation_date(analysis):
    """
    Extracts the domain creation date from VirusTotal URL analysis metadata.
    """
    try:
        attributes = analysis.get("attributes", {})
        whois = attributes.get("whois", "")
        # Parse the WHOIS data to extract the 'Creation Date' field
        for line in whois.splitlines():
            if "Creation Date" in line:
                creation_date = line.split(": ", 1)[-1].strip()
                logging.info(f"Extracted creation date: {creation_date}")
                return creation_date
    except Exception as e:  # pylint: disable=W0718
        logging.error(f"Error extracting creation date: {e}")
    return None

def _is_recent_domain(creation_date, recent_threshold_days):
    """Checks if the domain creation date is within the recent threshold."""
    try:
        creation_datetime = datetime.strptime(creation_date, "%Y-%m-%dT%H:%M:%S%z")
        days_since_creation = (datetime.now(creation_datetime.tzinfo) - creation_datetime).days
        logging.info(f"Days since domain creation: {days_since_creation}")
        return days_since_creation <= recent_threshold_days
    except ValueError as e:
        logging.error(f"Error parsing creation date '{creation_date}': {e}")
        return False

def check_email_for_attachment(file_path):
    """
    Checks if an email contains an attachment.
    """
    attachments = []
    try:
        with open(file_path, 'rb') as f:
            msg = email.message_from_binary_file(f, policy=policy.default)

        for part in msg.walk():
            if part.get_content_disposition() == 'attachment':
                filename = part.get_filename()
                if filename:
                    filepath = _save_attachment(part, filename)
                    if filepath:
                        attachments.append(filepath)
    except Exception as e:  # pylint: disable=W0718
        logging.error(f"Error parsing EML file: {e}")
    return attachments

def _save_attachment(part, filename):
    """
    Saves an attachment to a temporary file.
    """
    try:
        folder = "extracted_attachments"
        os.makedirs(folder, exist_ok=True)
        filepath = os.path.join(folder, filename)

        with open(filepath, 'wb') as f:
            f.write(part.get_payload(decode=True))

        logging.info(f"Saved attachment: {filepath}")
        return filepath
    except Exception as e:  # pylint: disable=W0718
        logging.error(f"Error saving attachment: {e}")
        return None

def hash_file(file_path):
    """
    Hashes a file using MD5.
    """
    with open(file_path, "rb") as f:
        file_hash = hashlib.md5(f.read()).hexdigest()
    return file_hash

def check_file_hash_in_virustotal(file_hash):
    """
    Checkc the file hash for analysis in Virustotal
    """
    with vt.Client(VIRUSTOTAL_API_KEY) as client:
        try:
            analysis = client.get_object(f"/files/{file_hash}")
            stats = analysis.last_analysis_stats
            malicious_count = stats.get('malicious', 0)

            logging.info(f"Hash analysis results: {stats}")
            return malicious_count
        except Exception as e:  # pylint: disable=W0718
            logging.exception(f"Unexpected error while querying VirusTotal: {e}")
            return None

def submit_file_to_virustotal(file_path):
    """
    Submits a file to Virustotal for analysis
    """
    with vt.Client(VIRUSTOTAL_API_KEY) as client:
        try:
            with open(file_path, "rb") as f:
                analysis = client.scan_file(f,wait_for_completion=True)
                analysis_id = analysis.id
                logging.info(f"File '{os.path.basename(file_path)}' submitted successfully with ID: {analysis_id}")

            analysis_results = analysis.stats
            logging.info(f"Analysis completed with stats: {analysis.stats}")
            
            if analysis_results['malicious'] > 0:
                logging.warning(f"File '{os.path.basename(file_path)}' is flagged as malicious!")
                return analysis_results['malicious']
            else:
                logging.info(f"File '{os.path.basename(file_path)}' appears clean.")
                return 0

        except Exception as e:  # pylint: disable=W0718
            logging.exception(f"Unexpected error while submitting file: {e}")

def parse_email_for_qrcode(file_path):
    """
    Parse the email file for a QR code.
    """
    detected_urls = []
    try:
        with open(file_path, 'rb') as f:
            msg = email.message_from_binary_file(f, policy=policy.default)

        # Check for QR codes in the email body
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type().startswith('image/'):
                    img_data = part.get_payload(decode=True)
                    urls = _extract_qr_code_urls(img_data)
                    detected_urls.extend(urls)

        # Check for QR codes in attachments
        for part in msg.walk():
            if part.get_content_disposition() == 'attachment':
                filename = part.get_filename()
                if filename:
                    attachment_data = part.get_payload(decode=True)
                    if filename.endswith(('png', 'jpg', 'jpeg')):
                        urls = _extract_qr_code_urls(attachment_data)
                        detected_urls.extend(urls)
                    elif filename.endswith('.pdf'):
                        urls = _extract_qr_code_from_pdf(attachment_data)
                        detected_urls.extend(urls)
                    elif filename.endswith(('doc', 'docx')):
                        urls = _extract_qr_code_from_docx(attachment_data)
                        detected_urls.extend(urls)
    except Exception as e:  # pylint: disable=W0718
        logging.error(f"Error parsing EML file: {e}")

    return detected_urls

def _extract_qr_code_urls(img_data):
    urls = []
    try:
        image = Image.open(BytesIO(img_data))
        decoded_objects = decode(image)
        for obj in decoded_objects:
            qr_data = obj.data.decode('utf-8')
            if qr_data.startswith('http'):
                logging.info(f"QR Code detected with URL: {qr_data}")
                urls.append(qr_data)
    except Exception as e:  # pylint: disable=W0718
        logging.error(f"Failed to decode QR code: {e}")

    return urls

def _extract_qr_code_from_pdf(pdf_data):
    """Extracts QR codes from PDF attachments."""
    urls = []
    try:
        pdf = fitz.open(stream=pdf_data, filetype="pdf")
        for page_num in range(pdf.page_count):
            page = pdf.load_page(page_num)
            images = page.get_images(full=True)
            for _, img in enumerate(images):
                xref = img[0]
                base_image = pdf.extract_image(xref)
                image_bytes = base_image["image"]
                urls.extend(_extract_qr_code_urls(image_bytes))
    except Exception as e:  # pylint: disable=W0718
        logging.error(f"Failed to extract QR code from PDF: {e}")

    return urls

def _extract_qr_code_from_docx(doc_data):
    """Extracts QR codes from DOCX attachments."""
    urls = []
    try:
        with BytesIO(doc_data) as doc_stream:
            doc = Document(doc_stream)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    img_data = rel.target_part.blob
                    urls.extend(_extract_qr_code_urls(img_data))
    except Exception as e:  # pylint: disable=W0718
        logging.error(f"Failed to extract QR code from DOCX: {e}")

    return urls

def _detect_fake_invoice(email_body, sender_email):
    """Detect if the email mentions a brand but the sender domain doesn't match."""
    sender_email = sender_email.split('<')[-1].strip().rstrip('>')
    email_domain = sender_email.split('@')[-1].lower()
    for brand in KNOWN_BRANDS:
        if brand in email_body.lower() and brand not in email_domain:
            logging.warning(f"Possible Fake Invoice scam detected: Brand '{brand}' mentioned in email, but sent from {email_domain}")
            return f"Possible Fake Invoice scam detected: Brand '{brand}' mentioned in email, but sent from {email_domain}"
    logging.info("Possible Fake Invoice scam not detected.")
    return "Possible Fake Invoice scam not detected."

def _extract_pdf_text(pdf_data):
    """Extract text from a PDF attachment."""
    try:
        pdf = fitz.open(stream=pdf_data, filetype="pdf")
        text = ""
        for page_num in range(pdf.page_count):
            page = pdf.load_page(page_num)
            text += page.get_text()
        return text
    except Exception as e:  # pylint: disable=W0718
        logging.error(f"Error extracting text from PDF: {e}")
        return ""

def _extract_docx_text(doc_data):
    """Extract text from a DOCX attachment."""
    try:
        with BytesIO(doc_data) as doc_stream:
            doc = Document(doc_stream)
            return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:  # pylint: disable=W0718
        logging.error(f"Error extracting text from DOCX: {e}")
        return ""

def check_sender_spoofing(sender_email):
    """Check for sender spoofing by comparing the sender's name and email address."""
    name, email = parseaddr(sender_email)
    
    # Normalize by lowercasing both the name and email for comparison
    name_parts = re.findall(r'\w+', name.lower())  # Extract individual words from the name
    email_username = email.split('@')[0].lower()   # Extract username part of the email
    
    name_in_email = any(part in email_username for part in name_parts)
    
    return name_in_email

def _find_phishing_indicators(text):
    """Detect phishing keywords and pattern matches in the provided text."""
    detected_keywords = [kw for kw in PHISHING_KEYWORDS if kw in text.lower()]
    pattern_matches = {name: re.findall(pattern, text) for name, pattern in PATTERNS.items()}
    if detected_keywords:
        logging.info(f"Possible Phishing indicators found: {detected_keywords}")
    else:
        logging.info("No phishing indicators found for the specific set of keywords.")
    return detected_keywords
def get_email_and_attachment_text(msg):
    """Extract email body and attachment text from a parsed email message."""
    fake_invoice_found = ''
    phishing_indicators_found = []

    for part in msg.walk():
        if part.get_content_disposition() == 'attachment':
            filename = part.get_filename()
            logging.info(f"Found attachment: {filename}")

            # Process PDF attachments
            if filename.endswith('.pdf'):
                pdf_data = part.get_payload(decode=True)
                pdf_text = _extract_pdf_text(pdf_data)
                fake_invoice_found = _detect_fake_invoice(pdf_text, msg.get('From'))
                phishing_indicators_found = _find_phishing_indicators(pdf_text)

            elif filename.endswith(('.docx', '.doc')):
                doc_data = part.get_payload(decode=True)
                doc_text = _extract_docx_text(doc_data)
                fake_invoice_found = _detect_fake_invoice(doc_text, msg.get('From'))
                phishing_indicators_found = _find_phishing_indicators(doc_text)


    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == 'text/plain':
                email_text = part.get_payload(decode=True).decode()
                fake_invoice_found = _detect_fake_invoice(email_text, msg.get('From'))
                phishing_indicators_found = _find_phishing_indicators(email_text)
    else:
        email_text = msg.get_payload(decode=True).decode()
        fake_invoice_found = _detect_fake_invoice(email_text, msg.get('From'))
        phishing_indicators_found = _find_phishing_indicators(email_text)
    return phishing_indicators_found, fake_invoice_found

def generate_final_report(phishing_indicators_found, fake_invoice_found, name_in_email,
                    qr_url_analysis, attachment_analysis, url_analysis, sender_analysis,
                    header_anomalies, dkim_status, dmarc_status, spf_status, is_malicious):
    """
    Generates a final report summarizing the findings of the email analysis.
    """
    report = """
# Email Analysis Report
## Warning
This report is not 100% accurate and should only be used as a guideline.
It is recommended to manually review the email and its contents to ensure accuracy.

## Summary\n"""

    report += "  ###Sender address Validation results\n"
    report += f"        - {spf_status}\n"
    report += f"        - {dmarc_status}\n"
    report += f"        - {dkim_status}\n\n"

    report += "  ###Sender email analysis\n"
    report += f"        - {sender_analysis}\n\n"

    report += "  ###Header Analysis\n"
    if header_anomalies:
        report += "     The following anomalies were found in the email headers: \n"
        for anomaly in header_anomalies:
            report += f"            - {anomaly}\n"
    else:
        report += "     No anomalies were found in the email headers.\n"

    report += "\n  ###URL Analysis\n"
    if url_analysis:
        report += "     The following URLs were found in the email body and the results are: \n"
        for url, result in url_analysis.items():
            report += f"        - {url}: {result}\n"
    else:
        report += "     The URLs are safe or no URLs were found in the email body.\n"


    report += "\n  ###Attachment Hash Analysis\n"
    if is_malicious:
        report += "     The attachment's hash is considered malicious.\n"
    else:
        report += "     The attachment's hash is not considered malicious.\n"

    report += "\n  ###Attachment Analysis\n"
    if attachment_analysis:
        report += "     The following is the analysis for the attachments.\n"
        for attachment, verdict in attachment_analysis.items():
            report += f"        - {attachment}: malicious_score - {verdict}\n"
    else:
        report += "     The attachment is not considered malicious or no attachments are present in the email\n"

    report += "\n  ###QR Code URL Analysis\n"
    if qr_url_analysis:
        report += "     The QR Code URLs have been analyzed and the results are: \n"
        for url, result in qr_url_analysis.items():
            report += f"       - {url}: {result}\n"
    else:
        report += "     No QR Code URLs were found in the email body.\n"

    report += "\n  ###Sender Spoofing\n"
    if name_in_email:
        report += "     **No Sender Spoofing Detected**: The sender's email address appears to be legitimate.\n"
    else:
        report += "     **Possible Sender Spoofing Detected**: The sender's email address MAY be spoofed OR sender name is not present in email\n"

    report += "\n  ###Phishing Indicators\n"
    if phishing_indicators_found:
        report += "     **Phishing Indicators Found**: The email text contains indicators of phishing.\n"
    else:
        report += "     **No Phishing Indicators Found**: The email text does not contain indicators of phishing.\n"

    report += "\n  ###Fake Invoice\n"
    report += f"     {fake_invoice_found}\n"

    return report

# pylint:disable=missing-function-docstring
def main():
    print("Advanced Phishing Email Analyzer\n")

    # Guide the user to save the email as .eml
    how_to_save_email()

    # Get file path of the email
    file_path = input("Enter the full path of the saved email file including the file name: ")
    if not os.path.isfile(file_path):
        print("Error: The provided path does not point to a valid file.")
        sys.exit(1)

    msg = parse_email(file_path)
    if msg is None:
        sys.exit(1)

    # Validate sender
    sender_address = msg['from']
    spf_status, dmarc_status, dkim_status = validate_sender_address(email_from=sender_address, msg=msg)

    # Header analysis
    headers = extract_email_headers(msg)
    sender_analysis = analyze_sender_address(headers)
    header_anomalies = analyze_headers_for_anomalies(headers)

    # URL analysis
    urls = extract_urls(msg)
    url_analysis = check_url_with_virustotal(urls)

    # Attachment analysis
    attachments = check_email_for_attachment(file_path)
    attachment_analysis = {}
    is_malicious = None
    
    if not attachments:
        logging.info("No attachments found in the EML file.")

    else:
        for attachment in attachments:
            file_hash = hash_file(attachment)
            if not file_hash:
                logging.error(f"Could not calculate hash for '{attachment}'. Skipping...")
                continue

            is_malicious = check_file_hash_in_virustotal(file_hash)

            if not is_malicious:
                logging.info(f"Submitting '{attachment}' for dynamic analysis...")
                attachment_analysis[attachment] = submit_file_to_virustotal(attachment)

    # Qr code detection and analysis
    qr_urls = parse_email_for_qrcode(file_path)

    if not qr_urls:
        logging.info("No QR codes found in the email.")
        qr_url_analysis = None

    else:
        qr_url_analysis = check_url_with_virustotal(qr_urls)

    # Sender spoofing detection
    sender_email = msg.get('From')
    name_in_email = check_sender_spoofing(sender_email)

    # Email and attachment text analysis
    phishing_indicators_found, fake_invoice_found = get_email_and_attachment_text(msg)

    report = generate_final_report(phishing_indicators_found, fake_invoice_found, name_in_email,
                    qr_url_analysis, attachment_analysis, url_analysis, sender_analysis,
                    header_anomalies, dkim_status, dmarc_status, spf_status, is_malicious)

    report_file_path = "email_analysis_report.txt"
    if os.path.isfile(report_file_path):
        os.remove(report_file_path)
    
    with open(report_file_path, "w") as report_file:
        report_file.write(report)

    print(f"Report written to {report_file_path}")

if __name__ == "__main__":
    main()
